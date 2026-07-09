import json
import uuid
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.core.config import settings
from app.core.security import get_current_user, require_role
from app.models.document import Document
from app.models.laureat import Laureat
from app.schemas.document import DocumentOut
from app.services.cv_analysis_service import analyze_cv, apply_analysis_to_laureat
from app.services.matching_service import run_matching

router = APIRouter(prefix="/api/documents", tags=["Documents"])

ALLOWED_TYPES = {"CV", "DIPLOME", "CIN", "LETTRE_MOTIVATION", "AUTRE"}
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".jpg", ".jpeg", ".png"}


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        return ""


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        tables = [
            cell.text for table in doc.tables for row in table.rows for cell in row.cells if cell.text
        ]
        return "\n".join(paragraphs + tables)
    except Exception:
        return ""


def _extract_cv_text(path: Path, ext: str) -> str:
    if ext == ".pdf":
        return _extract_pdf_text(path)
    if ext == ".docx":
        return _extract_docx_text(path)
    return ""


@router.post("/upload", response_model=DocumentOut)
def upload_document(
    type: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(require_role("laureat")),
):
    if type not in ALLOWED_TYPES:
        raise HTTPException(400, f"Type de document invalide. Valeurs autorisées : {', '.join(ALLOWED_TYPES)}")

    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Extension non autorisée. Formats acceptés : {', '.join(ALLOWED_EXTENSIONS)}")

    id_laureat = current_user.id_laureat
    laureat = db.query(Laureat).filter(Laureat.id_laureat == id_laureat).first()
    if not laureat:
        raise HTTPException(404, "Profil lauréat introuvable")

    dest_dir = settings.UPLOAD_DIR / id_laureat
    dest_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{type}_{uuid.uuid4().hex[:8]}{ext}"
    dest_path = dest_dir / safe_name

    content = file.file.read()
    dest_path.write_bytes(content)

    matching_a_relancer = False

    if type == "CV" and ext in (".pdf", ".docx"):
        extracted = _extract_cv_text(dest_path, ext)
        if extracted.strip():
            laureat.cv_text = extracted

            analysis = analyze_cv(extracted)
            if analysis is not None:
                apply_analysis_to_laureat(laureat, analysis)
                laureat.cv_analyse_json = json.dumps(analysis, ensure_ascii=False)
                laureat.cv_analyse_statut = "ok"
                matching_a_relancer = True
            else:
                laureat.cv_analyse_statut = "desactivee" if not settings.GEMINI_API_KEY else "echec"
        else:
            # Extraction du texte impossible (mise en page non standard, fichier corrompu...) :
            # on le signale plutot que de laisser cv_analyse_statut a None en silence.
            laureat.cv_analyse_statut = "extraction_echouee"
        laureat.cv_file_path = str(dest_path)
        db.add(laureat)
    elif type == "CV" and ext == ".doc":
        # Format binaire legacy non supporte (necessiterait un outil externe type
        # antiword/LibreOffice). Le fichier est bien stocke, mais ni le texte ni
        # l'analyse IA ne sont extraits.
        laureat.cv_analyse_statut = "format_non_supporte"
        laureat.cv_file_path = str(dest_path)
        db.add(laureat)

    doc = Document(
        id_laureat=id_laureat,
        type=type,
        nom_fichier=file.filename or safe_name,
        chemin_fichier=str(dest_path),
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    if matching_a_relancer:
        run_matching(db, id_laureat=id_laureat)

    return doc


@router.get("/me", response_model=list[DocumentOut])
def my_documents(db: Session = Depends(get_db), current_user=Depends(require_role("laureat"))):
    return db.query(Document).filter(Document.id_laureat == current_user.id_laureat).all()


@router.get("/{document_id}/download")
def download_document(document_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(404, "Document non trouvé")
    if current_user.role != "admin" and current_user.id_laureat != doc.id_laureat:
        raise HTTPException(403, "Accès non autorisé à ce document")
    path = Path(doc.chemin_fichier)
    if not path.exists():
        raise HTTPException(404, "Fichier introuvable sur le serveur")
    return FileResponse(path, filename=doc.nom_fichier)
