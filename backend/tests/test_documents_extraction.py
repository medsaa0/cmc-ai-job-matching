from pathlib import Path
from docx import Document as DocxDocument
from app.routers.documents import _extract_docx_text, _extract_cv_text


def test_extract_docx_text_reads_paragraphs(tmp_path):
    path = tmp_path / "cv.docx"
    doc = DocxDocument()
    doc.add_paragraph("Karim Benali")
    doc.add_paragraph("Python, Django, PostgreSQL")
    doc.save(str(path))

    text = _extract_docx_text(Path(path))
    assert "Karim Benali" in text
    assert "Python, Django, PostgreSQL" in text


def test_extract_docx_text_reads_table_cells(tmp_path):
    path = tmp_path / "cv_table.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Competence"
    table.rows[0].cells[1].text = "React"
    doc.save(str(path))

    text = _extract_docx_text(Path(path))
    assert "React" in text


def test_extract_docx_text_invalid_file_returns_empty(tmp_path):
    path = tmp_path / "not_a_docx.docx"
    path.write_bytes(b"ceci n'est pas un docx valide")
    assert _extract_docx_text(Path(path)) == ""


def test_extract_cv_text_dispatches_by_extension(tmp_path):
    path = tmp_path / "cv.docx"
    doc = DocxDocument()
    doc.add_paragraph("Contenu docx")
    doc.save(str(path))

    assert "Contenu docx" in _extract_cv_text(Path(path), ".docx")
    assert _extract_cv_text(Path(path), ".doc") == ""
